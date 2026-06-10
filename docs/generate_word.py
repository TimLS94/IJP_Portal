import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SHOTS = os.path.join(os.path.dirname(__file__), "screenshots")
OUT   = os.path.join(os.path.dirname(__file__), "anleitung-unternehmen-jobon.docx")

BLUE  = RGBColor(0x1e, 0x40, 0xaf)
LBLUE = RGBColor(0x3b, 0x82, 0xf6)
GRAY  = RGBColor(0x6b, 0x72, 0x80)
DGRAY = RGBColor(0x37, 0x41, 0x51)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_left_border(cell, color="3B82F6"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:color"), color)
    tcBorders.append(left)
    tcPr.append(tcBorders)

def add_cover(doc):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, "1E40AF")

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("JobOn – Internationale Fachkr\xe4fte")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xBF, 0xDB, 0xFE)
    r.font.bold = True

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run("Benutzerhandbuch\nf\xfcr Unternehmen")
    r2.font.size = Pt(26)
    r2.font.bold = True
    r2.font.color.rgb = WHITE

    p3 = cell.add_paragraph()
    p3.paragraph_format.space_before = Pt(8)
    p3.paragraph_format.space_after = Pt(24)
    r3 = p3.add_run("Schritt-f\xfcr-Schritt Anleitung zur Nutzung des JobOn-Portals — von der Registrierung bis zur erfolgreichen Einstellung.")
    r3.font.size = Pt(11)
    r3.font.color.rgb = RGBColor(0xBF, 0xDB, 0xFE)

    doc.add_paragraph()
    p4 = doc.add_paragraph()
    r4 = p4.add_run("Version 1.0  \xb7  Juni 2026  \xb7  jobon.work")
    r4.font.size = Pt(9)
    r4.font.color.rgb = GRAY
    doc.add_page_break()

def add_toc(doc):
    h = doc.add_heading("Inhaltsverzeichnis", level=1)
    h.runs[0].font.color.rgb = BLUE
    for e in [
        "1.  Zugang & Erstregistrierung",
        "2.  Das Dashboard – \xdcbersicht",
        "3.  Unternehmensprofil einrichten",
        "4.  Stellenanzeigen erstellen & verwalten",
        "5.  Bewerbungen verwalten",
        "6.  Vorstellungsgespr\xe4che planen",
        "7.  Team-Mitglieder verwalten",
        "8.  Einstellungen & Benachrichtigungen",
        "9.  H\xe4ufige Fragen (FAQ)",
    ]:
        p = doc.add_paragraph(e)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.runs[0].font.size = Pt(11)
        p.runs[0].font.color.rgb = LBLUE
    doc.add_paragraph()

def chapter_heading(doc, num, title):
    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=2)
    c0 = tbl.cell(0, 0)
    c0.width = Cm(1.5)
    set_cell_bg(c0, "1E40AF")
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(4)
    p0.paragraph_format.space_after = Pt(4)
    r0 = p0.add_run(str(num))
    r0.font.size = Pt(16)
    r0.font.bold = True
    r0.font.color.rgb = WHITE
    c1 = tbl.cell(0, 1)
    set_cell_bg(c1, "F0F4FF")
    p1 = c1.paragraphs[0]
    p1.paragraph_format.space_before = Pt(6)
    p1.paragraph_format.space_after = Pt(6)
    r1 = p1.add_run(title)
    r1.font.size = Pt(16)
    r1.font.bold = True
    r1.font.color.rgb = BLUE
    doc.add_paragraph()

def section_heading(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = BLUE
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:color"), "3B82F6")
    pBdr.append(left)
    pPr.append(pBdr)

def body_text(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    if p.runs:
        p.runs[0].font.size = Pt(11)
        p.runs[0].font.color.rgb = DGRAY

def bullet(doc, bold_part, rest):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(bold_part + " ")
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = DGRAY
    r2 = p.add_run(rest)
    r2.font.size = Pt(11)
    r2.font.color.rgb = DGRAY

def steps(doc, items):
    for i, (title, desc) in enumerate(items, 1):
        tbl = doc.add_table(rows=1, cols=2)
        c0 = tbl.cell(0, 0)
        c0.width = Cm(1.0)
        set_cell_bg(c0, "EFF6FF")
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_before = Pt(3)
        p0.paragraph_format.space_after = Pt(3)
        r0 = p0.add_run(str(i))
        r0.font.bold = True
        r0.font.size = Pt(11)
        r0.font.color.rgb = BLUE
        c1 = tbl.cell(0, 1)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(3)
        p1.paragraph_format.space_after = Pt(3)
        r_t = p1.add_run(title + "\n")
        r_t.font.bold = True
        r_t.font.size = Pt(11)
        r_t.font.color.rgb = DGRAY
        r_d = p1.add_run(desc)
        r_d.font.size = Pt(10)
        r_d.font.color.rgb = GRAY
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(2)

def note(doc, text, kind="info"):
    colors = {"info": ("EFF6FF", "1E40AF", "3B82F6"), "warn": ("FFFBEB", "92400E", "F59E0B"), "tip": ("F0FDF4", "166534", "22C55E")}
    bg, fg, border = colors.get(kind, colors["info"])
    icons = {"info": "ℹ️  ", "warn": "⚠️  ", "tip": "\U0001f4a1  "}
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    set_left_border(cell, border)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(icons.get(kind, "") + text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(fg)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)

def add_table(doc, headers, rows_data):
    tbl = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        set_cell_bg(cell, "F0F4FF")
        r = cell.paragraphs[0].add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = BLUE
    for ri, row in enumerate(rows_data, 1):
        for ci, val in enumerate(row):
            r = tbl.cell(ri, ci).paragraphs[0].add_run(val)
            r.font.size = Pt(10)
            r.font.color.rgb = DGRAY
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)

def add_image(doc, filename, caption):
    img_path = os.path.join(SHOTS, filename)
    if not os.path.exists(img_path):
        body_text(doc, "[Screenshot: " + caption + "]")
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img_path, width=Cm(15.5))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = GRAY
    cap.runs[0].font.italic = True

# ─── Dokument ────────────────────────────────────────────────────────────────
doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

add_cover(doc)
add_toc(doc)

# ─── Kapitel 1 ───────────────────────────────────────────────────────────────
chapter_heading(doc, 1, "Zugang & Erstregistrierung")
section_heading(doc, "Registrierung")
body_text(doc, "Sie k\xf6nnen sich auf zwei Wegen registrieren:")
bullet(doc, "Mit Einladungslink –", "JobOn schickt Ihnen einen Link. Ihr Konto ist sofort freigeschaltet.")
bullet(doc, "Ohne Einladungslink –", "Registrierung direkt auf jobon.work. Ihr Konto wird manuell gepr\xfcft und freigeschaltet (1–2 Werktage).")

section_heading(doc, "Registrierung mit Einladungslink")
steps(doc, [
    ("Einladungslink \xf6ffnen", "Klicken Sie auf den Link von JobOn. Sie werden zur Registrierungsseite weitergeleitet."),
    ("Formulardaten ausf\xfcllen", "Unternehmensname, Ansprechpartner, E-Mail-Adresse und Passwort."),
    ("Konto best\xe4tigen", "Sie erhalten eine Best\xe4tigungs-E-Mail – klicken Sie auf den Link darin."),
    ("Anmelden", "Rufen Sie jobon.work auf und melden Sie sich an."),
])
note(doc, "F\xfcr einen Einladungslink schreiben Sie uns: business@jobon.work", "info")

section_heading(doc, "Passwort vergessen")
steps(doc, [
    ("Passwort vergessen klicken", "Den Link finden Sie unterhalb des Anmeldeformulars."),
    ("E-Mail-Adresse eingeben", "Sie erhalten eine Reset-E-Mail (g\xfcltig 24 Stunden)."),
    ("Neues Passwort festlegen", "Klicken Sie den Link in der E-Mail und vergeben Sie ein neues Passwort."),
])

# ─── Kapitel 2 ───────────────────────────────────────────────────────────────
chapter_heading(doc, 2, "Das Dashboard – Ihre \xdcbersicht")
body_text(doc, "Nach dem Anmelden landen Sie automatisch auf Ihrem Dashboard.")
add_image(doc, "dashboard.png", "Dashboard – Kennzahlen und neueste Bewerbungen")
section_heading(doc, "Die vier Kennzahlen-Karten")
add_table(doc, ["Karte", "Was wird gezeigt?"], [
    ["Stellenanzeigen gesamt", "Alle je erstellten Stellenanzeigen"],
    ["Aktive Stellen", "Derzeit ver\xf6ffentlichte Anzeigen"],
    ["Bewerbungen erhalten", "Gesamtanzahl eingegangener Bewerbungen"],
    ["Offene Bewerbungen", "Bewerbungen im Status Ausstehend"],
])
section_heading(doc, "Neueste Bewerbungen")
body_text(doc, "Die letzten 5 Bewerbungen werden direkt angezeigt. Ein Klick \xf6ffnet die Details.")
note(doc, "Das Dashboard aktualisiert sich automatisch.", "tip")

# ─── Kapitel 3 ───────────────────────────────────────────────────────────────
chapter_heading(doc, 3, "Unternehmensprofil einrichten")
body_text(doc, "Navigieren Sie zu: Men\xfc → Unternehmensprofil")
add_image(doc, "profile.png", "Unternehmensprofil – Logo, Firmendaten, Beschreibung")
section_heading(doc, "Logo hochladen")
steps(doc, [
    ("Auf das Logo-Feld klicken", "Klicken Sie auf den Upload-Bereich oder ziehen Sie eine Bilddatei hinein."),
    ("Datei ausw\xe4hlen", "Formate: JPG, PNG, WebP. Max. 5 MB."),
    ("Speichern", "Klicken Sie auf Änderungen speichern."),
])
section_heading(doc, "Unternehmensdaten ausf\xfcllen")
add_table(doc, ["Feld", "Hinweis"], [
    ["Unternehmensname", "Vollst\xe4ndiger offizieller Name"],
    ["Ansprechpartner", "Zust\xe4ndige Person f\xfcr JobOn-Kommunikation"],
    ["Telefonnummer", "Nur f\xfcr JobOn intern sichtbar"],
    ["Website", "Optionaler Link zu Ihrer Website"],
    ["Branche", "Z. B. Gastronomie, Landwirtschaft, Produktion"],
    ["Unternehmensgr\xf6\xdfe", "Anzahl der Mitarbeiter"],
    ["Unternehmensvorstellung", "Erscheint auf Ihren Stellenanzeigen"],
])
note(doc, "Vergessen Sie nicht, auf Änderungen speichern zu klicken.", "warn")

# ─── Kapitel 4 ───────────────────────────────────────────────────────────────
chapter_heading(doc, 4, "Stellenanzeigen erstellen & verwalten")
body_text(doc, "Navigieren Sie zu: Men\xfc → Stellen")
section_heading(doc, "Neue Stellenanzeige erstellen")
steps(doc, [
    ("Neue Stelle klicken", "Schaltfl\xe4che oben rechts auf der Stellen\xfcbersicht."),
    ("Grunddaten eingeben", "Stellentitel, Standort, Einsatzdaten (Von – Bis), Positionstyp."),
    ("Stellenbeschreibung verfassen", "Je detaillierter, desto bessere Matches erhalten Sie."),
    ("Sprachanforderungen festlegen", "Erforderliches Deutschniveau (A1–C2) und optionale Englischkenntnisse."),
    ("Ver\xf6ffentlichen oder als Entwurf speichern", "Ver\xf6ffentlichen macht die Stelle sofort sichtbar."),
])
add_image(doc, "jobs.png", "Stellenanzeigen – aktive Stellen mit Metriken")
section_heading(doc, "Stellenanzeigen-Status")
add_table(doc, ["Status", "Bedeutung"], [
    ["Aktiv", "\xd6ffentlich sichtbar, Bewerber k\xf6nnen sich bewerben"],
    ["Entwurf", "Noch nicht ver\xf6ffentlicht, nur f\xfcr Sie sichtbar"],
    ["Inaktiv", "Tempor\xe4r deaktiviert, jederzeit wieder aktivierbar"],
    ["Archiviert", "Besetzt oder abgelaufen"],
])
note(doc, "JobOn \xfcbersetzt Ihre Stellenanzeigen automatisch in weitere Sprachen.", "info")

# ─── Kapitel 5 ───────────────────────────────────────────────────────────────
chapter_heading(doc, 5, "Bewerbungen verwalten")
body_text(doc, "Navigieren Sie zu: Men\xfc → Bewerbungen")
add_image(doc, "applications.png", "Bewerbungs\xfcbersicht – Statusfilter und Kandidatenliste")
section_heading(doc, "Der Bewerbungs-Workflow")
add_table(doc, ["Status", "Wann setzen?"], [
    ["⏳ Ausstehend", "Neue Bewerbung — wird automatisch gesetzt"],
    ["\U0001f535 In Pr\xfcfung", "Sie pr\xfcfen den Kandidaten aktiv"],
    ["\U0001f7e3 Gespr\xe4ch geplant", "Ein Termin wurde vorgeschlagen oder best\xe4tigt"],
    ["✅ Angenommen", "Kandidat wird eingestellt — JobOn wird benachrichtigt"],
    ["❌ Abgelehnt", "Kandidat passt nicht — automatische Absage-E-Mail"],
])
section_heading(doc, "Status einer Bewerbung \xe4ndern")
steps(doc, [
    ("Bewerbung finden", "Nutzen Sie Suche oder Filter (Status, Stelle, Match-Score)."),
    ("Status-Dropdown \xf6ffnen", "Klicken Sie direkt in der Tabelle auf den aktuellen Status-Badge."),
    ("Neuen Status w\xe4hlen", "Die \xc4nderung wird sofort gespeichert und der Bewerber per E-Mail benachrichtigt."),
])
section_heading(doc, "Kandidatendetails & Dokumente")
body_text(doc, "Klicken Sie auf den Namen eines Bewerbers um das Detaildialogfeld zu \xf6ffnen: vollst\xe4ndiges Profil, Kontaktdaten, Dokumente, Match-Score und Interview-Status.")
section_heading(doc, "Interne Notizen")
body_text(doc, "Das Notizen-Feld ist ausschlie\xdflich f\xfcr Ihr Team sichtbar — weder Bewerber noch JobOn k\xf6nnen es lesen.")
note(doc, "Auto-Ablehnung: \xdcber Einstellungen k\xf6nnen Sie einen Mindest-Match-Score festlegen. Gefilterte Bewerbungen bleiben erhalten.", "tip")

# ─── Kapitel 6 ───────────────────────────────────────────────────────────────
chapter_heading(doc, 6, "Vorstellungsgespr\xe4che planen")
body_text(doc, "Gespr\xe4che werden direkt aus dem Bewerbungsdetail heraus geplant. Der Kalender dient nur als \xdcbersicht.")
section_heading(doc, "Schritt 1 – Bewerbungsdetail \xf6ffnen")
body_text(doc, "Navigieren Sie zu Bewerbungen und klicken Sie auf das Auge-Symbol oder den Namen des Kandidaten. Scrollen Sie im Modal nach unten bis zum Abschnitt Vorstellungsgespr\xe4ch.")
add_image(doc, "application_detail.png", "Bewerbungsdetail – Kontakt, Profil, Dokumente und Interview-Bereich")
section_heading(doc, "Schritt 2 – Termine vorschlagen")
steps(doc, [
    ("Termine vorschlagen klicken", "Der Termin-Dialog \xf6ffnet sich. Sie k\xf6nnen 1–2 Alternativtermine vorschlagen."),
    ("Termin 1 ausf\xfcllen (Pflicht)", "Datum und Uhrzeit des ersten Vorschlags eingeben."),
    ("Termin 2 optional erg\xe4nzen", "Falls Sie eine Alternative anbieten m\xf6chten."),
    ("Ort oder Meeting-Link angeben", "F\xfcr Pr\xe4senz: Adresse. F\xfcr Video: Zoom-, Teams- oder Meet-Link."),
    ("Termine eintragen klicken", "Der Bewerber erh\xe4lt sofort eine E-Mail mit Ihren Terminvorschl\xe4gen."),
])
add_image(doc, "interview_scheduling.png", "Termine vorschlagen – bis zu 2 Alternativtermine")
section_heading(doc, "Schritt 3 – Antwort des Bewerbers")
add_table(doc, ["Symbol", "Bedeutung"], [
    ["✓  Best\xe4tigt", "Der Bewerber hat einen Termin angenommen."],
    ["✗  Abgelehnt", "Alle Vorschl\xe4ge abgelehnt. Bitte neue Termine vorschlagen."],
    ["⏳ Ausstehend", "Bewerber hat noch nicht geantwortet."],
])
note(doc, "Nach einer Best\xe4tigung \xe4ndert sich der Bewerbungsstatus automatisch auf Gespr\xe4ch geplant.", "info")
section_heading(doc, "Kalender – Gesamt\xfcbersicht")
body_text(doc, "Unter Men\xfc → Kalender sehen Sie alle geplanten Gespr\xe4che als Liste oder Monatsansicht. Best\xe4tigte Termine k\xf6nnen als ICS-Datei exportiert werden.")
add_image(doc, "calendar.png", "Kalender – alle geplanten Gespr\xe4che im \xdcberblick")

# ─── Kapitel 7 ───────────────────────────────────────────────────────────────
chapter_heading(doc, 7, "Team-Mitglieder verwalten")
body_text(doc, "Navigieren Sie zu: Men\xfc → Team")
add_image(doc, "team.png", "Teamverwaltung – Rollen und Mitglieder")
section_heading(doc, "Rollen & Berechtigungen")
add_table(doc, ["Rolle", "Berechtigungen"], [
    ["Eigent\xfcmer", "Vollzugriff — Stellen, Bewerbungen, Team, Einstellungen, Konto l\xf6schen"],
    ["Administrator", "Stellen und Bewerbungen verwalten + Team-Mitglieder einladen"],
    ["Mitglied", "Stellen bearbeiten und Bewerbungen verwalten — keine Teamverwaltung"],
])
section_heading(doc, "Neues Mitglied einladen")
steps(doc, [
    ("Mitglied einladen klicken", "Schaltfl\xe4che oben rechts auf der Team-Seite."),
    ("Daten eingeben", "E-Mail-Adresse, Vor- und Nachname sowie Rolle ausw\xe4hlen."),
    ("Einladung senden", "Die Person erh\xe4lt eine Einladungs-E-Mail."),
])
note(doc, "Der Eigent\xfcmer des Kontos kann nicht deaktiviert oder entfernt werden.", "warn")

# ─── Kapitel 8 ───────────────────────────────────────────────────────────────
chapter_heading(doc, 8, "Einstellungen & Benachrichtigungen")
body_text(doc, "Navigieren Sie zu: Men\xfc → Einstellungen")
add_image(doc, "settings.png", "Einstellungen – Benachrichtigungen und Sicherheit")
section_heading(doc, "E-Mail-Benachrichtigungen")
add_table(doc, ["Benachrichtigung", "Wann wird sie gesendet?"], [
    ["Neue Bewerbung", "Sofort wenn ein Bewerber sich bewirbt"],
    ["Bewerber-Digest", "T\xe4gliche Zusammenfassung aller neuen Bewerbungen"],
    ["Stellenanzeigen-Updates", "Bei wichtigen \xc4nderungen durch JobOn"],
])
section_heading(doc, "Automatische Absagen")
body_text(doc, "Unter Einstellungen → Auto-Ablehnung legen Sie einen Mindest-Match-Score fest. Den Absage-Text k\xf6nnen Sie individuell anpassen.")
section_heading(doc, "Passwort \xe4ndern")
steps(doc, [
    ("Einstellungen → Sicherheit → Passwort \xe4ndern", "Ein Dialogfeld \xf6ffnet sich."),
    ("Aktuelles und neues Passwort eingeben", "Das neue Passwort muss mindestens 8 Zeichen lang sein."),
    ("Best\xe4tigen", "Das neue Passwort ist sofort aktiv."),
])

# ─── Kapitel 9 ───────────────────────────────────────────────────────────────
chapter_heading(doc, 9, "H\xe4ufige Fragen (FAQ)")
faqs = [
    ("Warum sehe ich keine Bewerbungen?",
     "Die Stelle ist noch im Entwurf-Status, inaktiv, oder sehr neu."),
    ("Kann ich eine abgelehnte Bewerbung reaktivieren?",
     "Ja. \xd6ffnen Sie die Bewerbungsdetails und \xe4ndern Sie den Status zur\xfcck auf In Pr\xfcfung."),
    ("Wie lange sind Bewerbungen gespeichert?",
     "Bewerbungen bleiben dauerhaft erhalten — auch nach Archivierung der Stelle."),
    ("Erh\xe4lt der Bewerber automatisch eine E-Mail bei Status\xe4nderung?",
     "Ja — bei Gespr\xe4ch geplant, Angenommen und Abgelehnt wird automatisch eine E-Mail gesendet."),
    ("Was passiert wenn ich eine Stelle archiviere?",
     "Die Stelle verschwindet aus der \xf6ffentlichen Ansicht. Alle Bewerbungen bleiben erhalten."),
    ("Ich habe eine Frage die hier nicht beantwortet wird",
     "Schreiben Sie uns: business@jobon.work — wir antworten innerhalb von 1 Werktag."),
]
for q, a in faqs:
    section_heading(doc, q)
    body_text(doc, a)

# Footer
for _ in range(2):
    doc.add_paragraph()
p_f = doc.add_paragraph()
p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f = p_f.add_run("JobOn – Internationale Fachkr\xe4fte  \xb7  Husemannstr. 9, 10435 Berlin  \xb7  business@jobon.work  \xb7  jobon.work")
r_f.font.size = Pt(9)
r_f.font.color.rgb = GRAY
p_f2 = doc.add_paragraph()
p_f2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f2 = p_f2.add_run("Vertraulich – nur f\xfcr Partnerunternehmen von JobOn bestimmt.")
r_f2.font.size = Pt(9)
r_f2.font.italic = True
r_f2.font.color.rgb = GRAY

doc.save(OUT)
print("Word-Dokument gespeichert:", OUT)
