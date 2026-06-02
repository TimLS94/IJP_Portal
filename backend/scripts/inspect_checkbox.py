"""
Zeigt die XML-Struktur aller Zellen mit Checkbox-ähnlichen Inhalten in einer .docx-Datei.
Aufruf: python scripts/inspect_checkbox.py <datei.docx>
"""
import sys
from docx import Document
from lxml import etree

WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W14NS = "http://schemas.microsoft.com/office/word/2010/wordml"

CHECKBOX_HINTS = {"nein", "ні", "ja", "так", "geschlecht", "стать", "männlich", "weiblich"}


def cell_has_checkbox(cell):
    txt = cell.text.lower()
    return any(h in txt for h in CHECKBOX_HINTS)


def para_xml(para):
    return etree.tostring(para._element, pretty_print=True).decode()


def main(path):
    doc = Document(path)
    found = 0

    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if not cell_has_checkbox(cell):
                    continue
                print(f"\n{'='*70}")
                print(f"Tabelle {t_idx}, Zeile {r_idx}, Zelle {c_idx}")
                print(f"Text: {repr(cell.text[:120])}")
                print()
                for p_idx, para in enumerate(cell.paragraphs):
                    if not para.text.strip():
                        continue
                    print(f"  --- Absatz {p_idx}: {repr(para.text[:80])} ---")
                    # Runs
                    for ri, run in enumerate(para.runs):
                        print(f"    run[{ri}]: text={repr(run.text)} font={run.font.name}")
                    # w:sym
                    syms = para._element.findall(f".//{{{WNS}}}sym")
                    if syms:
                        for sym in syms:
                            print(f"    w:sym: {sym.attrib}")
                    # SDT (content controls)
                    sdts = para._element.findall(f".//{{{WNS}}}sdt")
                    if sdts:
                        print(f"    SDT-Elemente gefunden: {len(sdts)}")
                        for sdt in sdts:
                            print("    SDT XML:")
                            print("    " + etree.tostring(sdt, pretty_print=True).decode().replace("\n", "\n    ")[:500])
                    # fldChar
                    flds = para._element.findall(f".//{{{WNS}}}fldChar")
                    if flds:
                        print(f"    fldChar-Elemente: {len(flds)}")
                    # Raw XML (kompakt)
                    raw = etree.tostring(para._element).decode()
                    if any(c in raw for c in ["", "☐", "□", "sym", "sdt", "fldChar"]):
                        print(f"    RAW XML (gekürzt):")
                        print("    " + raw[:600])
                found += 1

    if found == 0:
        print("Keine Zellen mit Checkbox-Hinweisen gefunden.")
    else:
        print(f"\n{found} Zellen analysiert.")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Aufruf: python scripts/inspect_checkbox.py <datei.docx>")
        sys.exit(1)
    main(sys.argv[1])
