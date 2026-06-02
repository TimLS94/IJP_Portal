"""
IJP Dokumentenservice — Admin-Bereich
- CRUD für IJP-Betriebe
- Persistente, DB-gespeicherte Dokument-Vorlagen (editierbar)
- Template-basierte PDF-Generierung
"""
import io
import re
import os
import uuid
import logging
from datetime import datetime
from typing import Optional, List, Dict, Any
from pathlib import Path

from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

from app.core.database import get_db
from app.core.security import get_current_user
from app.core.config import settings
from app.models.user import User, UserRole
from app.models.ijp import IJPBetrieb, IJPTemplate, CRMContact, CompanyDocument
from app.models.applicant import Applicant

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/ijp", tags=["IJP"])


# ── Auth helper ────────────────────────────────────────────────────────────────

def _require_admin(current_user: User = Depends(get_current_user)) -> User:
    if current_user.role != UserRole.ADMIN:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Nur Admins")
    return current_user


# ── Seed-Daten (werden beim ersten Start in die DB geschrieben) ────────────────

SEED_TEMPLATES = [
    {
        "doc_type": "vollmacht",
        "label": "Vollmacht",
        "template_text": """\
Vollmacht

Bevollmächtigender (Arbeitgeber):
{{betrieb_name}}
{{contact_person}}
{{street}}
{{postal_code}} {{city}}

Bevollmächtigter:
IJP International Job Placement UG
Tim Schäfer
Lichtenberger Straße 34
10179 Berlin


Gegenstand der Vollmacht:

Hiermit bevollmächtigt {{betrieb_bezeichnung}} {{betrieb_name}}, vertreten durch {{contact_person}}, die Firma IJP International Job Placement UG mit Sitz in Berlin, die folgenden Aufgaben im Zusammenhang mit der Beantragung von Arbeitsgenehmigungen und Vorabzustimmungen zu übernehmen:

1. Erstellung der erforderlichen Unterlagen für die Beantragung von Arbeitsgenehmigungen bei der Bundesagentur für Arbeit.

2. Einreichung der Unterlagen bei den zuständigen Behörden.

3. Kommunikation mit der Bundesagentur für Arbeit sowie anderen relevanten Behörden in Bezug auf die Arbeitsgenehmigungen und Vorabzustimmungen.

4. Empfang der erteilten Arbeitsgenehmigungen, die an die Anschrift von IJP International Job Placement gesendet werden sollen:
IJP International Job Placement, c/o Schäfer, Husemannstr. 9, 10435 Berlin.

Diese Vollmacht gilt bis auf Widerruf und kann jederzeit schriftlich von {{betrieb_name}}, vertreten durch {{contact_person}}, widerrufen werden.


[[SPALTEN_START]]
{{city}}, den {{date}}
Unterschrift des Bevollmächtigenden




{{contact_person}}
({{betrieb_name}})
[[SPALTEN_MITTE]]
Berlin, den {{date}}
Unterschrift des Bevollmächtigten




Tim Schäfer
(IJP International Job Placement UG)
[[SPALTEN_ENDE]]""",
        "variables": [
            {"key": "betrieb_name",        "label": "Firmenname"},
            {"key": "contact_person",      "label": "Ansprechpartner"},
            {"key": "street",              "label": "Straße"},
            {"key": "postal_code",         "label": "PLZ"},
            {"key": "city",                "label": "Ort"},
            {"key": "betrieb_bezeichnung", "label": "Bezeichnung (z.B. das Restaurant, der Betrieb, das Hotel)"},
            {"key": "date",                "label": "Datum (z.B. 01.06.2025)"},
        ],
    },
    {
        "doc_type": "wohnungsbestaetigung",
        "label": "Wohnungsbestätigung",
        "template_text": """\
{{betrieb_name}}
{{contact_person}}
{{street}}
{{postal_code}} {{city}}


An die
Deutsche Botschaft


Betreff: Bestätigung zur Unterkunft und Übernahme der Anreisekosten


Sehr geehrte Damen und Herren,

hiermit bestätigen wir, dass {{gender_article}} {{applicant_name}} für den gesamten Zeitraum {{gender_possessive}} Beschäftigung in unserem Betrieb eine Unterkunft erhält. Die Unterkunft wird von uns organisiert und während der Dauer des Arbeitsverhältnisses zur Verfügung gestellt.

Des Weiteren erklären wir, dass wir bereit sind, die Anreisekosten des Arbeitnehmers zu übernehmen.

Diese Bestätigung erfolgt zur Vorlage bei der deutschen Botschaft im Rahmen des Visumverfahrens.

Für Rückfragen stehen wir Ihnen gerne zur Verfügung.

Mit freundlichen Grüßen


{{contact_person}}""",
        "variables": [
            {"key": "betrieb_name",       "label": "Betriebsname"},
            {"key": "contact_person",     "label": "Ansprechpartner"},
            {"key": "street",             "label": "Straße"},
            {"key": "postal_code",        "label": "PLZ"},
            {"key": "city",               "label": "Ort"},
            {"key": "applicant_name",     "label": "Name des Bewerbers"},
            {"key": "gender_article",     "label": "Geschlechtsartikel"},
            {"key": "gender_possessive",  "label": "Possessivpronomen (ihrer/seiner)"},
        ],
    },
]


def seed_ijp_templates(db: Session) -> None:
    """Schreibt fehlende Seed-Vorlagen in die DB (einmalig)."""
    for seed in SEED_TEMPLATES:
        exists = db.query(IJPTemplate).filter(IJPTemplate.doc_type == seed["doc_type"]).first()
        if not exists:
            db.add(IJPTemplate(**seed))
    db.commit()


# ── Helper ────────────────────────────────────────────────────────────────────

def _substitute(template: str, variables: Dict[str, str]) -> str:
    for key, value in variables.items():
        template = template.replace(f"{{{{{key}}}}}", value)
    return template


# ── Schemas ───────────────────────────────────────────────────────────────────

class BetriebCreate(BaseModel):
    name: str
    contact_person: Optional[str] = None
    street: Optional[str] = None
    postal_code: Optional[str] = None
    city: Optional[str] = None
    country: Optional[str] = None
    betriebsnummer: Optional[str] = None
    phone: Optional[str] = None
    email: Optional[str] = None
    website: Optional[str] = None
    industry: Optional[str] = None
    status: Optional[str] = None
    notes: Optional[str] = None


class BetriebResponse(BaseModel):
    id: int
    name: str
    contact_person: Optional[str] = None
    street: Optional[str] = None
    postal_code: Optional[str] = None
    city: Optional[str] = None
    country: Optional[str] = None
    betriebsnummer: Optional[str] = None
    phone: Optional[str] = None
    email: Optional[str] = None
    website: Optional[str] = None
    industry: Optional[str] = None
    status: Optional[str] = None
    notes: Optional[str] = None

    class Config:
        from_attributes = True


# ── CRM Schemas ───────────────────────────────────────────────────────────────

class ContactCreate(BaseModel):
    first_name: Optional[str] = None
    last_name: Optional[str] = None
    salutation: Optional[str] = None
    title: Optional[str] = None
    department: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    mobile: Optional[str] = None
    is_primary: bool = False


class ContactResponse(BaseModel):
    id: int
    company_id: int
    first_name: Optional[str] = None
    last_name: Optional[str] = None
    salutation: Optional[str] = None
    title: Optional[str] = None
    department: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    mobile: Optional[str] = None
    is_primary: bool = False

    class Config:
        from_attributes = True


class CompanyWithContacts(BetriebResponse):
    contacts: List[ContactResponse] = []


class ApplicantOption(BaseModel):
    id: int
    first_name: str
    last_name: str
    email: str
    gender: Optional[str] = None


class TemplateVar(BaseModel):
    key: str
    label: str


class TemplateCreate(BaseModel):
    doc_type: str
    label: str
    template_text: str
    variables: List[TemplateVar] = []


class TemplateUpdate(BaseModel):
    label: str
    template_text: str
    variables: List[TemplateVar] = []


class TemplateResponse(BaseModel):
    id: int
    doc_type: str
    label: str
    template_text: str
    variables: List[Dict[str, str]] = []
    created_at: datetime
    updated_at: datetime

    class Config:
        from_attributes = True


class DocumentRequest(BaseModel):
    doc_type: str
    betrieb_id: int
    applicant_id: Optional[int] = None
    gender: Optional[str] = None          # "female" | "male"
    custom_template: Optional[str] = None  # wenn None → DB-Template


# ── Betriebe CRUD ─────────────────────────────────────────────────────────────

@router.get("/betriebe", response_model=List[BetriebResponse])
def list_betriebe(db: Session = Depends(get_db), current_user: User = Depends(_require_admin)):
    return db.query(IJPBetrieb).order_by(IJPBetrieb.name).all()


@router.post("/betriebe", response_model=BetriebResponse, status_code=201)
def create_betrieb(data: BetriebCreate, db: Session = Depends(get_db), current_user: User = Depends(_require_admin)):
    betrieb = IJPBetrieb(**data.model_dump())
    db.add(betrieb)
    db.commit()
    db.refresh(betrieb)
    return betrieb


@router.put("/betriebe/{betrieb_id}", response_model=BetriebResponse)
def update_betrieb(betrieb_id: int, data: BetriebCreate, db: Session = Depends(get_db), current_user: User = Depends(_require_admin)):
    betrieb = db.query(IJPBetrieb).filter(IJPBetrieb.id == betrieb_id).first()
    if not betrieb:
        raise HTTPException(status_code=404, detail="Betrieb nicht gefunden")
    for key, value in data.model_dump().items():
        setattr(betrieb, key, value)
    betrieb.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(betrieb)
    return betrieb


@router.delete("/betriebe/{betrieb_id}", status_code=204)
def delete_betrieb(betrieb_id: int, db: Session = Depends(get_db), current_user: User = Depends(_require_admin)):
    betrieb = db.query(IJPBetrieb).filter(IJPBetrieb.id == betrieb_id).first()
    if not betrieb:
        raise HTTPException(status_code=404, detail="Betrieb nicht gefunden")
    db.delete(betrieb)
    db.commit()


# ── Bewerber ──────────────────────────────────────────────────────────────────

@router.get("/applicants", response_model=List[ApplicantOption])
def list_applicants_for_ijp(search: Optional[str] = None, db: Session = Depends(get_db), current_user: User = Depends(_require_admin)):
    query = db.query(Applicant, User).join(User, User.id == Applicant.user_id).filter(User.is_active == True)
    if search:
        term = f"%{search}%"
        query = query.filter(
            (Applicant.first_name.ilike(term)) | (Applicant.last_name.ilike(term)) | (User.email.ilike(term))
        )
    rows = query.order_by(Applicant.last_name, Applicant.first_name).limit(100).all()
    return [
        ApplicantOption(id=a.id, first_name=a.first_name or "", last_name=a.last_name or "",
                        email=u.email, gender=a.gender.value if a.gender else None)
        for a, u in rows
    ]


# ── Templates CRUD ────────────────────────────────────────────────────────────

@router.get("/templates", response_model=List[TemplateResponse])
def list_templates(db: Session = Depends(get_db), current_user: User = Depends(_require_admin)):
    return db.query(IJPTemplate).order_by(IJPTemplate.label).all()


@router.get("/templates/{doc_type}", response_model=TemplateResponse)
def get_template(doc_type: str, db: Session = Depends(get_db), current_user: User = Depends(_require_admin)):
    tmpl = db.query(IJPTemplate).filter(IJPTemplate.doc_type == doc_type).first()
    if not tmpl:
        raise HTTPException(status_code=404, detail="Vorlage nicht gefunden")
    return tmpl


@router.post("/templates", response_model=TemplateResponse, status_code=201)
def create_template(data: TemplateCreate, db: Session = Depends(get_db), current_user: User = Depends(_require_admin)):
    # Slug validieren
    if not re.match(r'^[a-z0-9_-]+$', data.doc_type):
        raise HTTPException(status_code=400, detail="doc_type darf nur Kleinbuchstaben, Ziffern, _ und - enthalten")
    existing = db.query(IJPTemplate).filter(IJPTemplate.doc_type == data.doc_type).first()
    if existing:
        raise HTTPException(status_code=409, detail=f"Vorlage '{data.doc_type}' existiert bereits")
    tmpl = IJPTemplate(
        doc_type=data.doc_type,
        label=data.label,
        template_text=data.template_text,
        variables=[v.model_dump() for v in data.variables],
    )
    db.add(tmpl)
    db.commit()
    db.refresh(tmpl)
    return tmpl


@router.put("/templates/{doc_type}", response_model=TemplateResponse)
def update_template(doc_type: str, data: TemplateUpdate, db: Session = Depends(get_db), current_user: User = Depends(_require_admin)):
    tmpl = db.query(IJPTemplate).filter(IJPTemplate.doc_type == doc_type).first()
    if not tmpl:
        raise HTTPException(status_code=404, detail="Vorlage nicht gefunden")
    tmpl.label = data.label
    tmpl.template_text = data.template_text
    tmpl.variables = [v.model_dump() for v in data.variables]
    tmpl.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(tmpl)
    return tmpl


@router.delete("/templates/{doc_type}", status_code=204)
def delete_template(doc_type: str, db: Session = Depends(get_db), current_user: User = Depends(_require_admin)):
    tmpl = db.query(IJPTemplate).filter(IJPTemplate.doc_type == doc_type).first()
    if not tmpl:
        raise HTTPException(status_code=404, detail="Vorlage nicht gefunden")
    db.delete(tmpl)
    db.commit()


# ── PDF-Generierung ───────────────────────────────────────────────────────────

def _build_pdf_from_filled_text(filled_text: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib import colors

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2.5*cm, rightMargin=2.5*cm, topMargin=2.5*cm, bottomMargin=2.5*cm)
    styles = getSampleStyleSheet()
    normal = ParagraphStyle("ijp_normal", parent=styles["Normal"], fontName="Helvetica", fontSize=11, leading=16, spaceAfter=2, alignment=TA_LEFT)

    def _safe(text: str) -> str:
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def _lines_to_story(text: str) -> list:
        items = []
        consecutive_blanks = 0
        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                consecutive_blanks += 1
                if consecutive_blanks <= 2:
                    items.append(Spacer(1, 0.45 * cm))
            else:
                consecutive_blanks = 0
                safe = _safe(stripped)
                if stripped.startswith("Betreff:"):
                    items.append(Paragraph(f"<b>{safe}</b>", normal))
                else:
                    items.append(Paragraph(safe, normal))
        return items

    # Split text into segments; [[SPALTEN_START]]...[[SPALTEN_MITTE]]...[[SPALTEN_ENDE]] creates two-column tables
    story = []
    segments = re.split(r'\[\[SPALTEN_START\]\]|\[\[SPALTEN_MITTE\]\]|\[\[SPALTEN_ENDE\]\]', filled_text)

    if len(segments) == 4:
        # before, left_col, right_col, after
        story.extend(_lines_to_story(segments[0]))
        left_items = _lines_to_story(segments[1])
        right_items = _lines_to_story(segments[2])
        col_width = (A4[0] - 5*cm) / 2
        table = Table([[left_items, right_items]], colWidths=[col_width, col_width])
        table.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ]))
        story.append(table)
        story.extend(_lines_to_story(segments[3]))
    else:
        story.extend(_lines_to_story(filled_text))

    doc.build(story)
    return buf.getvalue()


@router.post("/documents/generate")
def generate_document(req: DocumentRequest, db: Session = Depends(get_db), current_user: User = Depends(_require_admin)):
    tmpl = db.query(IJPTemplate).filter(IJPTemplate.doc_type == req.doc_type).first()
    if not tmpl:
        raise HTTPException(status_code=404, detail=f"Vorlage '{req.doc_type}' nicht gefunden")

    betrieb = db.query(IJPBetrieb).filter(IJPBetrieb.id == req.betrieb_id).first()
    if not betrieb:
        raise HTTPException(status_code=404, detail="Betrieb nicht gefunden")

    applicant_name = ""
    if req.applicant_id:
        applicant = db.query(Applicant).filter(Applicant.id == req.applicant_id).first()
        if not applicant:
            raise HTTPException(status_code=404, detail="Bewerber nicht gefunden")
        applicant_name = f"{applicant.first_name or ''} {applicant.last_name or ''}".strip()
        gender = req.gender or (applicant.gender.value if applicant.gender else "male")
    else:
        gender = req.gender or "male"

    variables = {
        "betrieb_name":      betrieb.name,
        "contact_person":    betrieb.contact_person or "",
        "street":            betrieb.street or "",
        "postal_code":       betrieb.postal_code or "",
        "city":              betrieb.city or "",
        "applicant_name":    applicant_name,
        "gender_article":    "die Arbeitnehmerin" if gender == "female" else "der Arbeitnehmer",
        "gender_possessive": "ihrer" if gender == "female" else "seiner",
    }

    template_text = req.custom_template or tmpl.template_text
    filled = _substitute(template_text, variables)
    pdf_bytes = _build_pdf_from_filled_text(filled)

    safe_betrieb = betrieb.name.replace(" ", "_")
    filename = f"{tmpl.label}_{safe_betrieb}.pdf" if not applicant_name else f"{tmpl.label}_{applicant_name.replace(' ', '_')}.pdf"

    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── CRM: Companies ────────────────────────────────────────────────────────────

@router.get("/crm/companies", response_model=List[CompanyWithContacts])
def crm_list_companies(
    search: Optional[str] = None,
    industry: Optional[str] = None,
    status: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(_require_admin),
):
    q = db.query(IJPBetrieb)
    if search:
        term = f"%{search}%"
        contact_match = db.query(CRMContact.company_id).filter(
            CRMContact.first_name.ilike(term)
            | CRMContact.last_name.ilike(term)
        ).subquery()
        q = q.filter(
            IJPBetrieb.name.ilike(term)
            | IJPBetrieb.city.ilike(term)
            | IJPBetrieb.industry.ilike(term)
            | IJPBetrieb.contact_person.ilike(term)
            | IJPBetrieb.id.in_(contact_match)
        )
    if industry:
        q = q.filter(IJPBetrieb.industry == industry)
    if status:
        q = q.filter(IJPBetrieb.status == status)
    return q.order_by(IJPBetrieb.name).all()


@router.post("/crm/companies", response_model=CompanyWithContacts, status_code=201)
def crm_create_company(
    data: BetriebCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(_require_admin),
):
    company = IJPBetrieb(**data.model_dump())
    db.add(company)
    db.commit()
    db.refresh(company)
    return company


@router.get("/crm/companies/{company_id}", response_model=CompanyWithContacts)
def crm_get_company(
    company_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(_require_admin),
):
    company = db.query(IJPBetrieb).filter(IJPBetrieb.id == company_id).first()
    if not company:
        raise HTTPException(status_code=404, detail="Firma nicht gefunden")
    return company


@router.put("/crm/companies/{company_id}", response_model=CompanyWithContacts)
def crm_update_company(
    company_id: int,
    data: BetriebCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(_require_admin),
):
    company = db.query(IJPBetrieb).filter(IJPBetrieb.id == company_id).first()
    if not company:
        raise HTTPException(status_code=404, detail="Firma nicht gefunden")
    for key, value in data.model_dump().items():
        setattr(company, key, value)
    company.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(company)
    return company


@router.delete("/crm/companies/{company_id}", status_code=204)
def crm_delete_company(
    company_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(_require_admin),
):
    company = db.query(IJPBetrieb).filter(IJPBetrieb.id == company_id).first()
    if not company:
        raise HTTPException(status_code=404, detail="Firma nicht gefunden")
    db.delete(company)
    db.commit()


# ── CRM: Contacts ─────────────────────────────────────────────────────────────

@router.post("/crm/companies/{company_id}/contacts", response_model=ContactResponse, status_code=201)
def crm_create_contact(
    company_id: int,
    data: ContactCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(_require_admin),
):
    company = db.query(IJPBetrieb).filter(IJPBetrieb.id == company_id).first()
    if not company:
        raise HTTPException(status_code=404, detail="Firma nicht gefunden")
    contact = CRMContact(company_id=company_id, **data.model_dump())
    db.add(contact)
    # If this is marked primary, unset others
    if data.is_primary:
        db.query(CRMContact).filter(
            CRMContact.company_id == company_id, CRMContact.id != contact.id
        ).update({"is_primary": False})
    db.commit()
    db.refresh(contact)
    return contact


@router.put("/crm/contacts/{contact_id}", response_model=ContactResponse)
def crm_update_contact(
    contact_id: int,
    data: ContactCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(_require_admin),
):
    contact = db.query(CRMContact).filter(CRMContact.id == contact_id).first()
    if not contact:
        raise HTTPException(status_code=404, detail="Kontakt nicht gefunden")
    for key, value in data.model_dump().items():
        setattr(contact, key, value)
    if data.is_primary:
        db.query(CRMContact).filter(
            CRMContact.company_id == contact.company_id, CRMContact.id != contact_id
        ).update({"is_primary": False})
    contact.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(contact)
    return contact


@router.delete("/crm/contacts/{contact_id}", status_code=204)
def crm_delete_contact(
    contact_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(_require_admin),
):
    contact = db.query(CRMContact).filter(CRMContact.id == contact_id).first()
    if not contact:
        raise HTTPException(status_code=404, detail="Kontakt nicht gefunden")
    db.delete(contact)
    db.commit()


@router.get("/crm/meta")
def crm_meta(db: Session = Depends(get_db), current_user: User = Depends(_require_admin)):
    """Distinct industries and statuses for filter dropdowns."""
    industries = [
        r[0] for r in db.query(IJPBetrieb.industry).filter(IJPBetrieb.industry.isnot(None)).distinct().all()
    ]
    statuses = [
        r[0] for r in db.query(IJPBetrieb.status).filter(IJPBetrieb.status.isnot(None)).distinct().all()
    ]
    return {"industries": sorted(industries), "statuses": sorted(statuses)}


# ── CRM: Company Documents ────────────────────────────────────────────────────

class CompanyDocumentResponse(BaseModel):
    id: int
    company_id: int
    name: str
    original_filename: str
    created_at: datetime
    class Config:
        from_attributes = True


def _applicant_context(applicant: Applicant, db: Session) -> dict:
    from app.models.user import User as UserModel
    user = db.query(UserModel).filter(UserModel.id == applicant.user_id).first()
    dob = applicant.date_of_birth.strftime("%d.%m.%Y") if applicant.date_of_birth else ""
    sb_start = applicant.semester_break_start.strftime("%d.%m.%Y") if applicant.semester_break_start else ""
    sb_end = applicant.semester_break_end.strftime("%d.%m.%Y") if applicant.semester_break_end else ""
    gender_val = applicant.gender.value if applicant.gender else ""
    return {
        "VORNAME":             applicant.first_name or "",
        "NACHNAME":            applicant.last_name or "",
        "NAME":                f"{applicant.first_name or ''} {applicant.last_name or ''}".strip(),
        "GEBURTSDATUM":        dob,
        "GEBURTSORT":          applicant.place_of_birth or "",
        "NATIONALITAET":       _resolve_nationality(applicant.nationality or ""),
        "STRASSE":             f"{applicant.street or ''} {applicant.house_number or ''}".strip(),
        "PLZ":                 applicant.postal_code or "",
        "ORT":                 applicant.city or "",
        "LAND":                applicant.country or "",
        "TELEFON":             applicant.phone or "",
        "EMAIL":               user.email if user else "",
        "UNI_NAME":            applicant.university_name or "",
        "UNI_STRASSE":         f"{applicant.university_street or ''} {applicant.university_house_number or ''}".strip(),
        "UNI_PLZ":             applicant.university_postal_code or "",
        "UNI_ORT":             applicant.university_city or "",
        "UNI_LAND":            applicant.university_country or "",
        "STUDIENGANG":         applicant.field_of_study or "",
        "SEMESTER":            str(applicant.current_semester) if applicant.current_semester else "",
        "SEMESTERFERIEN_VON":  sb_start,
        "SEMESTERFERIEN_BIS":  sb_end,
        "GESCHLECHT":          "weiblich" if gender_val == "female" else ("männlich" if gender_val == "male" else ""),
    }


# keyword → (context_key, underline, exclude_if_contains)
# Deutsche Nationalitäten — Index entspricht dem Wert der CustomSelect-Dropdown im Frontend
_NATIONALITIES_DE = [
    'Afghanisch','Ägyptisch','Albanisch','Algerisch','Amerikanisch','Andorranisch','Angolanisch','Antiguanisch','Äquatorialguineisch','Argentinisch','Armenisch','Aserbaidschanisch','Äthiopisch','Australisch',
    'Bahamaisch','Bahrainisch','Bangladeschisch','Barbadisch','Belarussisch','Belgisch','Belizisch','Beninisch','Bhutanisch','Bolivianisch','Bosnisch-Herzegowinisch','Botsuanisch','Brasilianisch','Britisch','Bruneiisch','Bulgarisch','Burkinisch','Burundisch',
    'Chilenisch','Chinesisch','Costa-ricanisch','Dänisch','Deutsch','Dominikanisch','Dschibutisch',
    'Ecuadorianisch','Salvadorianisch','Eritreisch','Estnisch','Eswatinisch',
    'Fidschi','Finnisch','Französisch',
    'Gabunisch','Gambisch','Georgisch','Ghanaisch','Grenadisch','Griechisch','Guatemaltekisch','Guineisch','Guinea-bissauisch','Guyanisch',
    'Haitianisch','Honduranisch',
    'Indisch','Indonesisch','Irakisch','Iranisch','Irisch','Isländisch','Israelisch','Italienisch','Ivorisch',
    'Jamaikanisch','Japanisch','Jemenitisch','Jordanisch',
    'Kambodschanisch','Kamerunisch','Kanadisch','Kapverdisch','Kasachisch','Katarisch','Kenianisch','Kirgisisch','Kiribatisch','Kolumbianisch','Komorisch','Kongolesisch','Kosovarisch','Kroatisch','Kubanisch','Kuwaitisch',
    'Laotisch','Lesothisch','Lettisch','Libanesisch','Liberianisch','Libysch','Liechtensteinisch','Litauisch','Luxemburgisch',
    'Madagassisch','Malawisch','Malaysisch','Maledivisch','Malisch','Maltesisch','Marokkanisch','Marshallisch','Mauretanisch','Mauritisch','Mazedonisch','Mexikanisch','Mikronesisch','Moldauisch','Monegassisch','Mongolisch','Montenegrinisch','Mosambikanisch','Myanmarisch',
    'Namibisch','Nauruisch','Nepalesisch','Neuseeländisch','Nicaraguanisch','Niederländisch','Nigerianisch','Nigrisch','Nordkoreanisch','Norwegisch',
    'Omanisch','Österreichisch',
    'Pakistanisch','Palauisch','Palästinensisch','Panamaisch','Papua-neuguineisch','Paraguayisch','Peruanisch','Philippinisch','Polnisch','Portugiesisch',
    'Ruandisch','Rumänisch','Russisch',
    'Salomonisch','Sambisch','Samoanisch','San-marinesisch','São-toméisch','Saudi-arabisch','Schwedisch','Schweizerisch','Senegalesisch','Serbisch','Seychellisch','Sierra-leonisch','Simbabwisch','Singapurisch','Slowakisch','Slowenisch','Somalisch','Spanisch','Sri-lankisch','Südafrikanisch','Sudanesisch','Südkoreanisch','Südsudanesisch','Surinamisch','Syrisch',
    'Tadschikisch','Tansanisch','Thailändisch','Timoresisch','Togoisch','Tongaisch','Trinidadisch','Tschadisch','Tschechisch','Tunesisch','Türkisch','Turkmenisch','Tuvaluisch',
    'Ugandisch','Ukrainisch','Ungarisch','Uruguayisch','Usbekisch',
    'Vanuatuisch','Vatikanisch','Venezolanisch','Vietnamesisch','Zentralafrikanisch','Zyprisch',
]


def _resolve_nationality(raw: str) -> str:
    """Wandelt numerischen Nationalitäts-Index in deutschen Text um."""
    if not raw:
        return ""
    stripped = raw.strip()
    if stripped.lstrip("-").isdigit():
        idx = int(stripped)
        if 0 <= idx < len(_NATIONALITIES_DE):
            return _NATIONALITIES_DE[idx]
        return stripped
    return stripped


_KEYWORD_MAP = [
    ("Name, Anschrift der Einrichtung",  "UNI_FULL",       False, []),
    ("Schulferien/Semesterferien",       "SEMESTER_RANGE",  False, []),
    ("Ende des Studiums",               "SCHULENTLASSUNG", False, []),
    ("Seit / З якого",                  "SEIT_STUDIUM",    False, []),
    ("Staatsangehörigkeit",             "NATIONALITAET",   False, []),
    ("Geburtsname",                     "NACHNAME",        False, []),  # Geburtsname = Nachname
    ("Geburtsort",                      "GEBURTSORT",      False, []),
    ("Derzeitige Adresse",              "STRASSE",         False, []),
    ("Geburtsdatum",                    "GEBURTSDATUM",    False, []),
    ("Postleitzahl",                    "PLZ",             False, []),
    ("Wohnort",                         "ORT",             False, []),
    ("Name /",                          "NACHNAME",        False, ["geburt"]),
    ("Vorname",                         "VORNAME",         True,  []),
]

_DRV_SKIP_KW   = ["bestätigung", "підтвердження", "подтвердження", "erklärung", "заява", "unterschrift"]
_DRV_SCHULE_KW = ["schule", "hochschule", "universität", "studium", "bildungseinrichtung", "besuchen sie"]


def _is_drv_form(doc) -> bool:
    for p in doc.paragraphs:
        if "versicherungspflicht" in p.text.lower() or "versicherungsfreiheit" in p.text.lower():
            return True
    return False


def _remove_para(para) -> None:
    """Entfernt einen Absatz aus dem XML (Zelle muss min. 1 Absatz behalten)."""
    parent = para._element.getparent()
    # Nur entfernen wenn noch mindestens 1 anderer Absatz übrig bleibt
    siblings = parent.findall(para._element.tag)
    if len(siblings) > 1:
        parent.remove(para._element)
    else:
        # Letzten Absatz nur leeren, nicht entfernen
        for run in para.runs:
            run.text = ""


def _mark_nein_ja(cell, mark: str) -> None:
    """Gewählte Option mit 'X   ' markieren, alle anderen stehen lassen."""
    target = "nein" if mark == "nein" else "ja"
    for para in cell.paragraphs:
        pt = para.text.strip().lower()
        if pt.startswith(target) and para.runs:
            if not para.runs[0].text.startswith("X   "):
                para.runs[0].text = "X   " + para.runs[0].text
                para.runs[0].bold = True


def _apply_drv_nein_ja(doc) -> None:
    """Mark nein/ja checkboxes in DRV-style government forms."""
    if not _is_drv_form(doc):
        return

    seen_refs: list = []
    seen_ids: set = set()

    for table in doc.tables:
        header_text = " ".join(c.text.lower() for c in (table.rows[0].cells if table.rows else []))

        # Skip personal data / signature / confirmation tables
        if any(kw in header_text for kw in _DRV_SKIP_KW):
            continue
        if "name /" in header_text or "vorname" in header_text or "geburtsdatum" in header_text:
            continue

        is_schulbesuch = any(kw in header_text for kw in _DRV_SCHULE_KW)

        for row in table.rows:
            row_text = " ".join(c.text.lower() for c in row.cells)

            for cell in row.cells:
                elem = cell._element
                eid = id(elem)
                if eid in seen_ids:
                    continue
                seen_ids.add(eid)
                seen_refs.append(elem)

                if not cell.paragraphs:
                    continue

                first_text = cell.paragraphs[0].text.strip().lower()

                if first_text.startswith("nein"):
                    # Standard nein/ja cell (separate paragraphs)
                    if is_schulbesuch and ("ende des studiums" in row_text or "schulentlassung" in row_text):
                        continue  # date field, no checkbox here
                    _mark_nein_ja(cell, "ja" if is_schulbesuch else "nein")
                else:
                    # Inline format: "nein / ні\t☐ ja / так…" as a later paragraph
                    if is_schulbesuch:
                        continue
                    for para in cell.paragraphs:
                        if (para.runs
                                and para.runs[0].text.lower() == "nein"
                                and not para.runs[0].text.startswith("X")):
                            para.runs[0].text = "X   " + para.runs[0].text
                            para.runs[0].bold = True


def _smart_fill_docx_bytes(file_bytes: bytes, context: dict) -> bytes:
    """
    Füllt ein .docx automatisch aus (aus bytes).
    - Hat das Dokument {{Platzhalter}} → docxtpl
    - Andernfalls → Keyword-Scan aller Tabellenzellen, Daten als neuen Absatz einfügen
    """
    from docx import Document as DocxDocument

    source = io.BytesIO(file_bytes)

    # Prüfe ob Platzhalter im Dokument vorhanden sind
    probe = DocxDocument(io.BytesIO(file_bytes))
    all_text = " ".join(p.text for p in probe.paragraphs)
    for tbl in probe.tables:
        for row in tbl.rows:
            for cell in row.cells:
                all_text += cell.text
    del probe

    if "{{" in all_text:
        from docxtpl import DocxTemplate
        tpl = DocxTemplate(io.BytesIO(file_bytes))
        tpl.render(context)
        buf = io.BytesIO()
        tpl.save(buf)
        buf.seek(0)
        return buf.read()

    # ── Keyword-basierter Filler ──────────────────────────────────────────────
    from docx.shared import Pt
    doc = DocxDocument(io.BytesIO(file_bytes))

    # Erweiterte Kontextwerte
    XX = "XX.XX.XXXX"
    uni_lines = [x for x in [
        context.get("UNI_NAME"), context.get("UNI_STRASSE"),
        f"{context.get('UNI_PLZ','')} {context.get('UNI_ORT','')}".strip() or None,
        context.get("UNI_LAND"),
    ] if x]
    context["UNI_FULL"] = "\n".join(uni_lines)

    sf_von = context.get("SEMESTERFERIEN_VON") or XX
    sf_bis = context.get("SEMESTERFERIEN_BIS") or XX
    context["SEMESTER_RANGE"] = f"vom {sf_von} bis {sf_bis}"
    context["SCHULENTLASSUNG"] = context.get("SEMESTERFERIEN_BIS") or XX
    context["SEIT_STUDIUM"] = XX  # Immatrikulationsdatum nicht im System → Platzhalter

    gender = context.get("GESCHLECHT", "")

    seen_refs: list = []
    seen_ids: set = set()

    for table in doc.tables:
        # Tabellenkontext: ist das die Schulbesuch/Studium-Tabelle?
        table_header = " ".join(c.text.lower() for c in table.rows[0].cells) if table.rows else ""
        is_schulbesuch_table = any(kw in table_header for kw in _DRV_SCHULE_KW)

        for row in table.rows:
            for cell in row.cells:
                elem = cell._element
                eid = id(elem)
                if eid in seen_ids:
                    continue
                seen_ids.add(eid)
                seen_refs.append(elem)

                cell_text = cell.text.strip()
                first_line = cell_text.split("\n")[0].strip()

                # Geschlecht: nur die gewählte Option mit X markieren, alle anderen stehen lassen
                if "Geschlecht" in first_line or "Стать" in first_line:
                    if gender:
                        options = ["männlich", "weiblich", "divers"]
                        for para in cell.paragraphs:
                            pt = para.text.strip()
                            if not any(o in pt.lower() for o in options):
                                continue
                            if gender.lower() in pt.lower() and para.runs:
                                para.runs[0].text = "X   " + para.runs[0].text
                                para.runs[0].bold = True
                    continue

                # Keyword-Match → Wert als neuen Absatz mit leichtem Abstand
                for keyword, data_key, underline, excludes in _KEYWORD_MAP:
                    if keyword.lower() not in first_line.lower():
                        continue
                    if any(e in first_line.lower() for e in excludes):
                        continue
                    # "Seit" nur in der Schulbesuch-Tabelle füllen
                    if data_key == "SEIT_STUDIUM" and not is_schulbesuch_table:
                        continue
                    value = context.get(data_key, "")
                    if not value:
                        break
                    first_added = True
                    for line in str(value).split("\n"):
                        line = line.strip()
                        if not line:
                            continue
                        new_para = cell.add_paragraph(line)
                        if first_added:
                            new_para.paragraph_format.space_before = Pt(4)
                            first_added = False
                        if new_para.runs:
                            new_para.runs[0].bold = True
                            new_para.runs[0].underline = underline
                    break

    # ── DRV Nein/Ja-Checkboxen ───────────────────────────────────────────────
    _apply_drv_nein_ja(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@router.get("/crm/companies/{company_id}/documents", response_model=List[CompanyDocumentResponse])
def list_company_documents(
    company_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(_require_admin),
):
    company = db.query(IJPBetrieb).filter(IJPBetrieb.id == company_id).first()
    if not company:
        raise HTTPException(status_code=404, detail="Firma nicht gefunden")
    return company.company_documents


@router.post("/crm/companies/{company_id}/documents", response_model=CompanyDocumentResponse, status_code=201)
async def upload_company_document(
    company_id: int,
    file: UploadFile = File(...),
    name: str = Form(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(_require_admin),
):
    company = db.query(IJPBetrieb).filter(IJPBetrieb.id == company_id).first()
    if not company:
        raise HTTPException(status_code=404, detail="Firma nicht gefunden")
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Nur .docx-Dateien werden unterstützt")

    contents = await file.read()

    doc = CompanyDocument(
        company_id=company_id,
        name=name,
        original_filename=file.filename,
        file_content=contents,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    return doc


@router.get("/crm/companies/{company_id}/documents/{doc_id}/download")
def download_company_document(
    company_id: int,
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(_require_admin),
):
    doc = db.query(CompanyDocument).filter(
        CompanyDocument.id == doc_id,
        CompanyDocument.company_id == company_id,
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Dokument nicht gefunden")
    return StreamingResponse(
        io.BytesIO(doc.file_content),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{doc.original_filename}"'},
    )


@router.delete("/crm/companies/{company_id}/documents/{doc_id}", status_code=204)
def delete_company_document(
    company_id: int,
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(_require_admin),
):
    doc = db.query(CompanyDocument).filter(
        CompanyDocument.id == doc_id,
        CompanyDocument.company_id == company_id,
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Dokument nicht gefunden")
    try:
        Path(doc.file_path).unlink(missing_ok=True)
    except Exception:
        pass
    db.delete(doc)
    db.commit()


@router.post("/crm/companies/{company_id}/documents/{doc_id}/fill")
def fill_company_document(
    company_id: int,
    doc_id: int,
    applicant_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(_require_admin),
):
    from docxtpl import DocxTemplate

    doc = db.query(CompanyDocument).filter(
        CompanyDocument.id == doc_id,
        CompanyDocument.company_id == company_id,
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Dokument nicht gefunden")

    applicant = db.query(Applicant).filter(Applicant.id == applicant_id).first()
    if not applicant:
        raise HTTPException(status_code=404, detail="Bewerber nicht gefunden")

    context = _applicant_context(applicant, db)
    filled_bytes = _smart_fill_docx_bytes(doc.file_content, context)

    safe_name = f"{applicant.first_name}_{applicant.last_name}".replace(" ", "_")
    filename = f"{doc.name}_{safe_name}.docx"

    return StreamingResponse(
        io.BytesIO(filled_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── IJP-Tab: Alle Firmendokumente auflisten + ausfüllen ───────────────────────

@router.get("/employer-docs", response_model=List[CompanyDocumentResponse])
def list_all_employer_docs(
    db: Session = Depends(get_db),
    current_user: User = Depends(_require_admin),
):
    """Alle hochgeladenen Firmendokumente für den IJP-Dokumente-Tab."""
    return db.query(CompanyDocument).order_by(CompanyDocument.company_id, CompanyDocument.name).all()


@router.post("/employer-docs/{doc_id}/fill")
def fill_employer_doc(
    doc_id: int,
    applicant_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(_require_admin),
):
    doc = db.query(CompanyDocument).filter(CompanyDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Dokument nicht gefunden")

    applicant = db.query(Applicant).filter(Applicant.id == applicant_id).first()
    if not applicant:
        raise HTTPException(status_code=404, detail="Bewerber nicht gefunden")

    context = _applicant_context(applicant, db)
    filled_bytes = _smart_fill_docx_bytes(doc.file_content, context)

    safe_name = f"{applicant.first_name}_{applicant.last_name}".replace(" ", "_")
    filename = f"{doc.name}_{safe_name}.docx"

    return StreamingResponse(
        io.BytesIO(filled_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
